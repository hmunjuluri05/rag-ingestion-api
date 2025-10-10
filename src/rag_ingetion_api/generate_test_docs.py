"""
Generate sample test documents in various formats for testing the extraction comparison tool.
"""

import os
import sys
from pathlib import Path

# Set UTF-8 encoding for Windows console
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')


def generate_markdown_docs(output_dir: Path):
    """Generate sample Markdown documents"""

    # Simple Markdown
    simple_md = """# Sample Markdown Document

## Introduction

This is a sample markdown document created for testing document extraction capabilities.

## Features

Markdown supports various formatting options:

- **Bold text**
- *Italic text*
- Lists (like this one)
- Code blocks
- Links and more

## Code Example

```python
def hello_world():
    print("Hello, World!")
```

## Table

| Feature | Unstructured | Azure DI |
|---------|-------------|----------|
| Speed   | Fast        | Moderate |
| Cost    | Free        | Paid     |

## Conclusion

This document tests basic markdown extraction capabilities including headings, lists, code blocks, and tables.
"""

    # Technical Markdown
    technical_md = """# Technical Documentation: RAG Pipeline Architecture

## Executive Summary

Retrieval-Augmented Generation (RAG) systems combine the power of large language models with external knowledge retrieval to provide more accurate and contextual responses.

## Architecture Components

### 1. Document Ingestion

The ingestion pipeline consists of several stages:

1. **Extraction**: Convert documents to text
2. **Chunking**: Split text into manageable segments
3. **Embedding**: Generate vector representations
4. **Indexing**: Store in vector database

### 2. Retrieval Mechanism

```python
class VectorRetriever:
    def __init__(self, index, embedding_model):
        self.index = index
        self.embedding_model = embedding_model

    def retrieve(self, query: str, top_k: int = 5):
        query_vector = self.embedding_model.encode(query)
        results = self.index.search(query_vector, k=top_k)
        return results
```

### 3. Performance Metrics

| Metric | Target | Current |
|--------|--------|---------|
| Latency (p95) | < 200ms | 150ms |
| Throughput | 1000 req/s | 850 req/s |
| Accuracy | > 90% | 92% |

## Best Practices

- Use semantic chunking for better context preservation
- Implement hybrid search (keyword + vector)
- Cache frequently accessed embeddings
- Monitor retrieval quality metrics

## Deployment

### Kubernetes Configuration

```yaml
apiVersion: apps/v1
kind: Deployment
metadata:
  name: rag-service
spec:
  replicas: 3
  template:
    spec:
      containers:
      - name: rag-api
        image: rag-service:latest
        resources:
          limits:
            memory: "2Gi"
            cpu: "1000m"
```

## Conclusion

A well-designed RAG pipeline balances performance, accuracy, and cost-effectiveness.
"""

    (output_dir / "01_simple.md").write_text(simple_md, encoding='utf-8')
    (output_dir / "02_technical.md").write_text(technical_md, encoding='utf-8')
    print(f"✓ Generated 2 Markdown documents")


def generate_html_docs(output_dir: Path):
    """Generate sample HTML documents"""

    # Simple HTML
    simple_html = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>Sample HTML Document</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; }
        h1 { color: #333; }
        .highlight { background-color: yellow; }
        table { border-collapse: collapse; width: 100%; }
        th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
    </style>
</head>
<body>
    <h1>HTML Document Extraction Test</h1>

    <h2>Introduction</h2>
    <p>This HTML document contains various elements to test extraction capabilities:</p>

    <ul>
        <li>Headings and paragraphs</li>
        <li>Lists (ordered and unordered)</li>
        <li>Tables with data</li>
        <li>Styled text with <span class="highlight">highlighting</span></li>
    </ul>

    <h2>Sample Table</h2>
    <table>
        <thead>
            <tr>
                <th>Product</th>
                <th>Price</th>
                <th>Stock</th>
            </tr>
        </thead>
        <tbody>
            <tr>
                <td>Widget A</td>
                <td>$19.99</td>
                <td>150</td>
            </tr>
            <tr>
                <td>Widget B</td>
                <td>$29.99</td>
                <td>85</td>
            </tr>
            <tr>
                <td>Widget C</td>
                <td>$39.99</td>
                <td>42</td>
            </tr>
        </tbody>
    </table>

    <h2>Nested Content</h2>
    <div>
        <p>This section contains <strong>nested</strong> and <em>formatted</em> text.</p>
        <blockquote>
            "The best way to predict the future is to invent it." - Alan Kay
        </blockquote>
    </div>

    <h2>Conclusion</h2>
    <p>This document tests how well extraction tools handle HTML structure and styling.</p>
</body>
</html>
"""

    # Form-like HTML
    form_html = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>Invoice Document</title>
</head>
<body>
    <h1>INVOICE</h1>

    <div style="margin-bottom: 20px;">
        <h3>Invoice Details</h3>
        <p><strong>Invoice Number:</strong> INV-2025-001</p>
        <p><strong>Date:</strong> October 9, 2025</p>
        <p><strong>Due Date:</strong> November 9, 2025</p>
    </div>

    <div style="margin-bottom: 20px;">
        <h3>Billed To</h3>
        <p>Acme Corporation<br>
        123 Business Street<br>
        New York, NY 10001</p>
    </div>

    <table style="width: 100%; border-collapse: collapse;">
        <thead>
            <tr style="background-color: #f0f0f0;">
                <th style="border: 1px solid #ddd; padding: 8px;">Item</th>
                <th style="border: 1px solid #ddd; padding: 8px;">Quantity</th>
                <th style="border: 1px solid #ddd; padding: 8px;">Unit Price</th>
                <th style="border: 1px solid #ddd; padding: 8px;">Total</th>
            </tr>
        </thead>
        <tbody>
            <tr>
                <td style="border: 1px solid #ddd; padding: 8px;">Professional Services</td>
                <td style="border: 1px solid #ddd; padding: 8px;">40 hours</td>
                <td style="border: 1px solid #ddd; padding: 8px;">$150.00</td>
                <td style="border: 1px solid #ddd; padding: 8px;">$6,000.00</td>
            </tr>
            <tr>
                <td style="border: 1px solid #ddd; padding: 8px;">Cloud Infrastructure</td>
                <td style="border: 1px solid #ddd; padding: 8px;">1 month</td>
                <td style="border: 1px solid #ddd; padding: 8px;">$2,500.00</td>
                <td style="border: 1px solid #ddd; padding: 8px;">$2,500.00</td>
            </tr>
            <tr>
                <td style="border: 1px solid #ddd; padding: 8px;">Support License</td>
                <td style="border: 1px solid #ddd; padding: 8px;">1</td>
                <td style="border: 1px solid #ddd; padding: 8px;">$500.00</td>
                <td style="border: 1px solid #ddd; padding: 8px;">$500.00</td>
            </tr>
        </tbody>
        <tfoot>
            <tr style="background-color: #f9f9f9; font-weight: bold;">
                <td colspan="3" style="border: 1px solid #ddd; padding: 8px; text-align: right;">Subtotal:</td>
                <td style="border: 1px solid #ddd; padding: 8px;">$9,000.00</td>
            </tr>
            <tr>
                <td colspan="3" style="border: 1px solid #ddd; padding: 8px; text-align: right;">Tax (10%):</td>
                <td style="border: 1px solid #ddd; padding: 8px;">$900.00</td>
            </tr>
            <tr style="background-color: #e0e0e0; font-weight: bold;">
                <td colspan="3" style="border: 1px solid #ddd; padding: 8px; text-align: right;">TOTAL:</td>
                <td style="border: 1px solid #ddd; padding: 8px;">$9,900.00</td>
            </tr>
        </tfoot>
    </table>

    <div style="margin-top: 30px;">
        <h3>Payment Terms</h3>
        <p>Payment is due within 30 days. Please make checks payable to "Tech Solutions Inc."</p>
    </div>
</body>
</html>
"""

    (output_dir / "03_simple_page.html").write_text(simple_html, encoding='utf-8')
    (output_dir / "04_invoice.html").write_text(form_html, encoding='utf-8')
    print(f"✓ Generated 2 HTML documents")


def generate_text_docs(output_dir: Path):
    """Generate sample plain text documents"""

    # Simple text
    simple_txt = """PLAIN TEXT DOCUMENT - SIMPLE

This is a simple plain text document for testing extraction capabilities.

SECTION 1: INTRODUCTION
Plain text files are the most basic document format. They contain no formatting,
just raw text characters. This makes them easy to process but limits their
ability to convey structure and styling.

SECTION 2: ADVANTAGES
- Universal compatibility
- Small file size
- Easy to parse and process
- No proprietary formats

SECTION 3: DISADVANTAGES
- No formatting options
- No embedded images
- Limited structure representation
- No styling capabilities

SECTION 4: USE CASES
Plain text is ideal for:
1. Configuration files
2. Log files
3. Simple data storage
4. Code source files
5. Documentation

CONCLUSION
Despite its limitations, plain text remains a fundamental and widely-used
document format in computing.
"""

    # Log-like text
    log_txt = """APPLICATION LOG FILE
===================

[2025-10-09 10:15:23] INFO: Application started successfully
[2025-10-09 10:15:24] INFO: Loading configuration from config.yaml
[2025-10-09 10:15:25] INFO: Database connection established
[2025-10-09 10:15:26] INFO: Redis cache initialized
[2025-10-09 10:15:27] INFO: API server listening on port 8000

[2025-10-09 10:16:01] INFO: Received document extraction request
[2025-10-09 10:16:01] DEBUG: File: document.pdf, Size: 2.5 MB
[2025-10-09 10:16:02] INFO: Starting extraction with Unstructured library
[2025-10-09 10:16:03] DEBUG: Extracted 45 elements
[2025-10-09 10:16:03] INFO: Extraction completed in 1.23s

[2025-10-09 10:17:30] WARNING: High memory usage detected: 85%
[2025-10-09 10:17:31] INFO: Triggered garbage collection
[2025-10-09 10:17:32] INFO: Memory usage reduced to 62%

[2025-10-09 10:18:45] ERROR: Failed to connect to external API
[2025-10-09 10:18:45] ERROR: Connection timeout after 30s
[2025-10-09 10:18:46] INFO: Retrying with exponential backoff
[2025-10-09 10:18:50] INFO: Connection successful on retry attempt 2

[2025-10-09 10:20:00] INFO: Processing batch of 100 documents
[2025-10-09 10:20:15] INFO: Batch processing: 25% complete
[2025-10-09 10:20:30] INFO: Batch processing: 50% complete
[2025-10-09 10:20:45] INFO: Batch processing: 75% complete
[2025-10-09 10:21:00] INFO: Batch processing: 100% complete
[2025-10-09 10:21:01] INFO: Total processing time: 61s
[2025-10-09 10:21:01] INFO: Average time per document: 0.61s

STATISTICS SUMMARY
------------------
Total Requests: 157
Successful: 152
Failed: 5
Success Rate: 96.8%
Average Response Time: 1.2s
Peak Memory Usage: 85%
"""

    (output_dir / "05_simple.txt").write_text(simple_txt, encoding='utf-8')
    (output_dir / "06_log_file.txt").write_text(log_txt, encoding='utf-8')
    print(f"✓ Generated 2 TXT documents")


def generate_docx_docs(output_dir: Path):
    """Generate sample DOCX documents using python-docx"""

    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        # Simple DOCX
        doc1 = Document()

        # Title
        title = doc1.add_heading('Sample DOCX Document', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Introduction
        doc1.add_heading('Introduction', 1)
        doc1.add_paragraph(
            'This is a sample Microsoft Word document created using python-docx. '
            'It demonstrates various formatting capabilities including headings, '
            'paragraphs, lists, and tables.'
        )

        # Lists
        doc1.add_heading('Key Features', 2)
        doc1.add_paragraph('Rich text formatting', style='List Bullet')
        doc1.add_paragraph('Multiple heading levels', style='List Bullet')
        doc1.add_paragraph('Tables and images', style='List Bullet')
        doc1.add_paragraph('Styled paragraphs', style='List Bullet')

        # Table
        doc1.add_heading('Sample Data Table', 2)
        table = doc1.add_table(rows=4, cols=3)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Category'
        header_cells[1].text = 'Value'
        header_cells[2].text = 'Status'

        # Data rows
        data = [
            ('Performance', '95%', 'Good'),
            ('Reliability', '98%', 'Excellent'),
            ('Efficiency', '87%', 'Good')
        ]

        for i, (cat, val, status) in enumerate(data, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = cat
            row_cells[1].text = val
            row_cells[2].text = status

        # Conclusion
        doc1.add_heading('Conclusion', 1)
        doc1.add_paragraph(
            'This document serves as a test case for document extraction tools, '
            'allowing comparison of how different libraries handle DOCX format.'
        )

        doc1.save(str(output_dir / '07_simple.docx'))

        # Technical Report DOCX
        doc2 = Document()

        # Title page
        title = doc2.add_heading('RAG System Performance Report', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        date_para = doc2.add_paragraph('October 9, 2025')
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc2.add_page_break()

        # Executive Summary
        doc2.add_heading('Executive Summary', 1)
        doc2.add_paragraph(
            'This report presents the performance analysis of our Retrieval-Augmented '
            'Generation (RAG) system over the past quarter. Key findings include improved '
            'response accuracy, reduced latency, and enhanced scalability.'
        )

        # Metrics
        doc2.add_heading('Performance Metrics', 1)
        doc2.add_heading('Latency Analysis', 2)
        doc2.add_paragraph(
            'Average response time: 150ms (target: 200ms)\n'
            'P95 latency: 280ms\n'
            'P99 latency: 450ms'
        )

        doc2.add_heading('Accuracy Metrics', 2)
        doc2.add_paragraph(
            'Document retrieval accuracy: 92%\n'
            'Response relevance score: 4.3/5.0\n'
            'User satisfaction: 88%'
        )

        # Comparison Table
        doc2.add_heading('Library Comparison', 1)
        table2 = doc2.add_table(rows=4, cols=4)
        table2.style = 'Medium Grid 1 Accent 1'

        headers = ['Library', 'Speed', 'Accuracy', 'Cost']
        for i, header in enumerate(headers):
            table2.rows[0].cells[i].text = header

        comparison_data = [
            ('Unstructured', 'Fast', 'Good', 'Free'),
            ('Azure DI', 'Moderate', 'Excellent', 'Paid'),
            ('PyPDF2', 'Very Fast', 'Fair', 'Free')
        ]

        for i, (lib, speed, acc, cost) in enumerate(comparison_data, 1):
            cells = table2.rows[i].cells
            cells[0].text = lib
            cells[1].text = speed
            cells[2].text = acc
            cells[3].text = cost

        # Recommendations
        doc2.add_heading('Recommendations', 1)
        doc2.add_paragraph('Based on our analysis, we recommend:', style='List Number')
        doc2.add_paragraph('Implementing hybrid extraction strategy', style='List Number')
        doc2.add_paragraph('Scaling worker pods during peak hours', style='List Number')
        doc2.add_paragraph('Adding semantic caching layer', style='List Number')

        doc2.save(str(output_dir / '08_technical_report.docx'))

        print(f"✓ Generated 2 DOCX documents")
        return True

    except ImportError:
        print(f"⚠ Skipping DOCX generation (python-docx not installed)")
        print(f"  Install with: pip install python-docx")
        return False


def generate_pdf_docs(output_dir: Path):
    """Generate sample PDF documents using reportlab"""

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib import colors

        # Simple PDF
        pdf1_path = output_dir / '09_simple.pdf'
        doc1 = SimpleDocTemplate(str(pdf1_path), pagesize=letter)
        story1 = []
        styles = getSampleStyleSheet()

        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f77b4'),
            spaceAfter=30,
            alignment=1  # Center
        )
        story1.append(Paragraph("Sample PDF Document", title_style))
        story1.append(Spacer(1, 12))

        # Introduction
        story1.append(Paragraph("Introduction", styles['Heading2']))
        story1.append(Paragraph(
            "This is a sample PDF document created using ReportLab. "
            "It demonstrates various PDF generation capabilities including text formatting, "
            "tables, and structured content.",
            styles['BodyText']
        ))
        story1.append(Spacer(1, 12))

        # Features list
        story1.append(Paragraph("Document Features", styles['Heading2']))
        features = [
            "Multi-level headings",
            "Formatted paragraphs with proper spacing",
            "Tables with styling",
            "Professional layout"
        ]
        for feature in features:
            story1.append(Paragraph(f"• {feature}", styles['BodyText']))
        story1.append(Spacer(1, 12))

        # Table
        story1.append(Paragraph("Sample Table", styles['Heading2']))
        table_data = [
            ['Feature', 'Supported', 'Notes'],
            ['Text Extraction', 'Yes', 'Full support'],
            ['Table Parsing', 'Yes', 'With structure'],
            ['Images', 'Partial', 'OCR required'],
        ]
        t1 = Table(table_data)
        t1.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story1.append(t1)
        story1.append(Spacer(1, 12))

        # Conclusion
        story1.append(Paragraph("Conclusion", styles['Heading2']))
        story1.append(Paragraph(
            "This PDF document provides a comprehensive test case for document extraction "
            "tools, allowing evaluation of text extraction, table parsing, and layout understanding.",
            styles['BodyText']
        ))

        doc1.build(story1)

        # Technical PDF with more complex content
        pdf2_path = output_dir / '10_technical.pdf'
        doc2 = SimpleDocTemplate(str(pdf2_path), pagesize=letter)
        story2 = []

        # Title Page
        story2.append(Paragraph("Document Extraction Benchmark Report", title_style))
        story2.append(Spacer(1, 0.5*inch))
        story2.append(Paragraph("Performance Analysis Q4 2025", styles['Heading3']))
        story2.append(PageBreak())

        # Executive Summary
        story2.append(Paragraph("Executive Summary", styles['Heading1']))
        story2.append(Paragraph(
            "This report presents a comprehensive analysis of document extraction libraries "
            "tested across multiple document formats and complexity levels. Our evaluation "
            "focused on extraction accuracy, processing speed, and resource utilization.",
            styles['BodyText']
        ))
        story2.append(Spacer(1, 12))

        # Test Results
        story2.append(Paragraph("Test Results", styles['Heading1']))
        story2.append(Paragraph("Performance Comparison", styles['Heading2']))

        results_data = [
            ['Library', 'Avg Time (s)', 'Success Rate', 'Text Quality'],
            ['Unstructured', '1.23', '98%', 'Good'],
            ['Azure DI', '2.45', '99%', 'Excellent'],
            ['PyPDF2', '0.87', '85%', 'Fair'],
            ['PyMuPDF', '0.95', '92%', 'Good'],
        ]

        t2 = Table(results_data, colWidths=[2*inch, 1.5*inch, 1.5*inch, 1.5*inch])
        t2.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.lightblue),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.lightblue, colors.white])
        ]))
        story2.append(t2)
        story2.append(Spacer(1, 12))

        # Recommendations
        story2.append(Paragraph("Recommendations", styles['Heading1']))
        recommendations = [
            "Use Unstructured for cost-sensitive applications with simple documents",
            "Deploy Azure DI for production systems requiring high accuracy",
            "Implement hybrid approach for optimal cost-performance balance",
            "Monitor extraction quality metrics continuously"
        ]
        for i, rec in enumerate(recommendations, 1):
            story2.append(Paragraph(f"{i}. {rec}", styles['BodyText']))
            story2.append(Spacer(1, 6))

        # Conclusion
        story2.append(Spacer(1, 12))
        story2.append(Paragraph("Conclusion", styles['Heading1']))
        story2.append(Paragraph(
            "Our analysis demonstrates that no single library is optimal for all use cases. "
            "The choice depends on specific requirements including budget constraints, "
            "accuracy requirements, and document complexity. A hybrid approach combining "
            "multiple libraries based on document characteristics yields the best results.",
            styles['BodyText']
        ))

        doc2.build(story2)

        print(f"✓ Generated 2 PDF documents")
        return True

    except ImportError:
        print(f"⚠ Skipping PDF generation (reportlab not installed)")
        print(f"  Install with: pip install reportlab")
        return False


def main():
    """Generate all test documents"""
    output_dir = Path("C:/Users/hanum/projects/rag-ingetion-api/test_documents")
    output_dir.mkdir(exist_ok=True)

    print("\nGenerating test documents...")
    print("=" * 80)

    # Always generated (no dependencies)
    generate_markdown_docs(output_dir)
    generate_html_docs(output_dir)
    generate_text_docs(output_dir)

    # Require optional dependencies
    docx_generated = generate_docx_docs(output_dir)
    pdf_generated = generate_pdf_docs(output_dir)

    print("=" * 80)
    print(f"\n✓ Test documents generated in: {output_dir}")

    # List all generated files
    files = sorted(output_dir.glob("*"))
    print(f"\nGenerated {len(files)} files:")
    for f in files:
        size = f.stat().st_size
        print(f"  - {f.name} ({size:,} bytes)")

    # Installation hints
    if not (docx_generated and pdf_generated):
        print("\n" + "=" * 80)
        print("To generate all document types, install optional dependencies:")
        if not docx_generated:
            print("  pip install python-docx")
        if not pdf_generated:
            print("  pip install reportlab")
        print("=" * 80)

    print("\nYou can now run the comparison tool:")
    print(f'  python compare_extractors.py "{output_dir}"')


if __name__ == "__main__":
    main()
